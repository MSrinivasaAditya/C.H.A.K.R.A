import docx
from docx.shared import Pt, RGBColor
import re

filepath = r'C:\Users\msrin\.gemini\antigravity\brain\8a5556d6-53db-4740-87d5-295ddb1d3cc1\artifacts\vibesentinel_progress_writeup.md'
outpath = r'C:\Users\msrin\.gemini\antigravity\brain\8a5556d6-53db-4740-87d5-295ddb1d3cc1\artifacts\CHAKRA_Week1_Report.docx'

doc = docx.Document()

with open(filepath, 'r', encoding='utf-8') as f:
    lines = f.readlines()

for line in lines:
    line = line.strip()
    if not line:
        continue
    if line.startswith('# '):
        heading = doc.add_heading(level=0)
        run = heading.add_run(line[2:])
        run.bold = True
    elif line.startswith('## '):
        doc.add_heading(line[3:], level=1)
    elif line.startswith('### '):
        doc.add_heading(line[4:], level=2)
    elif line.startswith('* ') or line.startswith('- '):
        p = doc.add_paragraph(style='List Bullet')
        text = line[2:]
        if text.startswith('**') and '**' in text[2:]:
            end_idx = text.find('**', 2)
            bold_text = text[2:end_idx]
            normal_text = text[end_idx+2:]
            run = p.add_run(bold_text)
            run.bold = True
            p.add_run(normal_text)
        else:
            p.add_run(text)
    elif line == '***':
        p = doc.add_paragraph()
        run = p.add_run('---')
    else:
        p = doc.add_paragraph()
        text = line
        
        # handle inline bold
        parts = re.split(r'(\*\*.*?\*\*)', text)
        for part in parts:
            if part.startswith('**') and part.endswith('**'):
                run = p.add_run(part[2:-2])
                run.bold = True
            else:
                p.add_run(part)

doc.save(outpath)
print("Docx created successfully at:", outpath)
